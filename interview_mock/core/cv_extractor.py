"""
CV Extractor Module - Trích xuất thông tin từ CV sử dụng LLM
"""
from __future__ import annotations

import re
import json
import base64
import tempfile
import os
from io import BytesIO
from typing import Dict, Optional
from pathlib import Path

import PyPDF2
from PIL import Image
from docx import Document
from pydantic import ValidationError

# LLM imports
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.prompts import PromptTemplate
from langchain.output_parsers import PydanticOutputParser

from core.models import CVData, WorkExperience, Education, Project
from core.skills import SkillManager
from config import Config

class CVExtractor:
    """Trích xuất thông tin từ CV sử dụng Google Gemini LLM"""
    
    def __init__(self, api_key: Optional[str] = None):
        """Initialize CV extractor với LLM"""
        self.api_key = api_key or Config.GOOGLE_API_KEY
        
        # Initialize LLM
        self.llm = ChatGoogleGenerativeAI(
            model=Config.LLM_MODEL,
            api_key=self.api_key,
            temperature=Config.LLM_TEMPERATURE
        )
        
        # Initialize parser
        self.parser = PydanticOutputParser(pydantic_object=CVData)
        
        # Skill manager for skill extraction
        self.skill_manager = SkillManager()
    
    def extract_from_file(self, file_path: str) -> CVData:
        """
        Trích xuất CV từ file
        Supports: PDF, DOCX, JPG, PNG
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File không tồn tại: {file_path}")
        
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext not in Config.SUPPORTED_FILE_TYPES:
            raise ValueError(f"File type không được hỗ trợ: {file_ext}. Hỗ trợ: {Config.SUPPORTED_FILE_TYPES}")
        
        try:
            if file_ext == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                return self._extract_from_docx(file_path)
            elif file_ext in ['.jpg', '.jpeg', '.png']:
                return self._extract_from_image(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
                
        except Exception as e:
            print(f"Lỗi extract CV: {str(e)}")
            # Fallback to basic extraction
            return self._fallback_extraction(file_path)
    
    def extract_from_text(self, text: str) -> CVData:
        """Trích xuất CV từ plain text"""
        if not text or not text.strip():
            raise ValueError("Text content trống")
        
        try:
            return self._process_with_llm(text)
        except Exception as e:
            print(f"Lỗi extract từ text: {str(e)}")
            return self._fallback_text_extraction(text)
    
    def _extract_from_pdf(self, pdf_path: str) -> CVData:
        """Extract từ PDF file"""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                
                if not text.strip():
                    raise ValueError("Không thể extract text từ PDF")
                
                return self._process_with_llm(text)
                
        except Exception as e:
            raise Exception(f"Lỗi đọc PDF: {str(e)}")
    
    def _extract_from_docx(self, docx_path: str) -> CVData:
        """Extract từ DOCX file"""
        try:
            doc = Document(docx_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            if not text.strip():
                raise ValueError("DOCX file trống")
            
            return self._process_with_llm(text)
            
        except Exception as e:
            raise Exception(f"Lỗi đọc DOCX: {str(e)}")
    
    def _extract_from_image(self, image_path: str) -> CVData:
        """Extract từ image file sử dụng Gemini Vision"""
        try:
            # Process image
            img_base64 = self._process_image(image_path)
            
            # Vision prompt
            vision_prompt = self._create_vision_prompt()
            
            # Call Gemini Vision
            messages = [
                {"type": "text", "text": vision_prompt},
                {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{img_base64}"}}
            ]
            
            response = self.llm.invoke(messages)
            
            # Parse response
            return self.parser.parse(response.content)
            
        except Exception as e:
            print(f"Lỗi extract từ image: {str(e)}")
            # Fallback to basic info
            return CVData()
    
    def _process_image(self, image_path: str) -> str:
        """Process image to base64"""
        try:
            image = Image.open(image_path)
            
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Resize if too large
            max_size = 1024
            if max(image.size) > max_size:
                image.thumbnail((max_size, max_size))
            
            # Convert to base64
            buffered = BytesIO()
            image.save(buffered, format="JPEG", quality=95)
            img_str = base64.b64encode(buffered.getvalue()).decode()
            
            return img_str
            
        except Exception as e:
            raise Exception(f"Lỗi xử lý image: {str(e)}")
    
    def _process_with_llm(self, text: str) -> CVData:
        """Process text với LLM để extract structured data"""
        
        # Tạo prompt
        prompt_template = self._create_extraction_prompt()
        
        prompt = PromptTemplate(
            template=prompt_template,
            input_variables=["cv_text"],
            partial_variables={"format_instructions": self.parser.get_format_instructions()}
        )
        
        # Create chain
        chain = prompt | self.llm | self.parser
        
        try:
            # Invoke chain
            result = chain.invoke({"cv_text": text})
            
            # Post-process result
            return self._post_process_cv_data(result)
            
        except Exception as e:
            print(f"LLM processing error: {str(e)}")
            # Fallback
            return self._fallback_text_extraction(text)
    
    def _create_extraction_prompt(self) -> str:
        """Tạo prompt để extract CV information"""
        return """
Bạn là một chuyên gia phân tích CV. Hãy extract thông tin từ CV sau thành format JSON structured.

HƯỚNG DẪN EXTRACT:

1. THÔNG TIN CÁ NHÂN:
- full_name: Tên đầy đủ
- email: Địa chỉ email
- phone: Số điện thoại
- location: Địa chỉ/Thành phố
- date_of_birth: Ngày sinh (nếu có)

2. THÔNG TIN NGHỀ NGHIỆP:
- current_position: Vị trí hiện tại hoặc mong muốn
- career_objective: Mục tiêu nghề nghiệp
- summary: Tóm tắt về bản thân

3. KỸ NĂNG:
- technical_skills: Các kỹ năng kỹ thuật (programming languages, frameworks, tools...)
- soft_skills: Kỹ năng mềm (teamwork, communication, leadership...)
- languages: Ngôn ngữ (English, Vietnamese... kèm level nếu có)

4. KINH NGHIỆM LÀM VIỆC:
- work_experience: Danh sách công việc với format:
  {{
    "position": "Vị trí công việc",
    "company": "Tên công ty", 
    "duration": "Thời gian",
    "description": "Mô tả công việc",
    "start_date": "Ngày bắt đầu",
    "end_date": "Ngày kết thúc",
    "is_current": true/false
  }}
- total_experience_years: Tổng số năm kinh nghiệm (số thập phân)

5. HỌC VẤN:
- education: Danh sách bằng cấp với format:
  {{
    "degree": "Bằng cấp",
    "university": "Trường học",
    "major": "Chuyên ngành",
    "year": năm_tốt_nghiệp,
    "gpa": điểm_GPA
  }}
- highest_education: Trình độ cao nhất (Trung học/Trung cấp/Cao đẳng/Đại học/Thạc sĩ/Tiến sĩ)

6. DỰ ÁN & CHỨNG CHỈ:
- projects: Danh sách dự án với format:
  {{
    "name": "Tên dự án",
    "description": "Mô tả dự án", 
    "technologies": ["tech1", "tech2"],
    "url": "Link dự án",
    "start_date": "Ngày bắt đầu",
    "end_date": "Ngày kết thúc"
  }}
- certifications: Danh sách chứng chỉ

LƯU Ý:
- Trích xuất chính xác, không bịa đặt thông tin
- Chuẩn hóa tên skills theo standard (ví dụ: js -> JavaScript, react -> ReactJS)
- Tính toán total_experience_years dựa trên work_experience
- Nếu không có thông tin, để null hoặc empty array
- Đảm bảo format JSON hợp lệ

CV TEXT:
{cv_text}

{format_instructions}

Hãy trả về JSON format chính xác theo schema:
"""
    
    def _create_vision_prompt(self) -> str:
        """Tạo prompt cho Gemini Vision"""
        return """
Bạn là một chuyên gia phân tích CV. Hãy phân tích ảnh CV này và extract thông tin theo định dạng JSON yêu cầu.

Hãy trích xuất các thông tin sau từ ảnh CV:

1. THÔNG TIN CÁ NHÂN: Tên, email, điện thoại, địa chỉ
2. THÔNG TIN NGHỀ NGHIỆP: Vị trí hiện tại, mục tiêu nghề nghiệp  
3. KỸ NĂNG: Technical skills, soft skills, ngôn ngữ
4. KINH NGHIỆM: Danh sách công việc với chi tiết
5. HỌC VẤN: Bằng cấp, trường học, chuyên ngành
6. DỰ ÁN & CHỨNG CHỈ: Các dự án và chứng chỉ

Lưu ý:
- Đọc kỹ ảnh và extract chính xác
- Không bịa đặt thông tin không có
- Chuẩn hóa tên skills
- Trả về JSON format hợp lệ

Hãy trả về kết quả theo đúng schema JSON đã định.
"""
    
    def _post_process_cv_data(self, cv_data: CVData) -> CVData:
        """Post-process CV data sau khi extract"""
        
        # Chuẩn hóa skills
        if cv_data.technical_skills:
            cv_data.technical_skills = self.skill_manager.validate_skills(cv_data.technical_skills)
        
        # Extract additional skills from text fields
        additional_skills = []
        
        # Extract from work experience descriptions
        for exp in cv_data.work_experience:
            if exp.description:
                skills = self.skill_manager.extract_skills_from_text(exp.description)
                additional_skills.extend(skills)
        
        # Extract skills from projects if available
        if cv_data.projects:
            for project in cv_data.projects:
                if project.technologies:
                    for tech in project.technologies:
                        tech_skills = self.skill_manager.extract_skills_from_text(tech)
                        additional_skills.extend(tech_skills)
                if project.description:
                    desc_skills = self.skill_manager.extract_skills_from_text(project.description)
                    additional_skills.extend(desc_skills)
        
        # Merge and deduplicate skills
        all_skills = set(cv_data.technical_skills + additional_skills)
        cv_data.technical_skills = list(all_skills)
        
        # Total experience years removed - use years_experience field instead
        
        return cv_data
    
    # Removed total experience calculation - using years_experience field directly
    
    def _fallback_extraction(self, file_path: str) -> CVData:
        """Fallback extraction using basic methods"""
        try:
            # Try to extract text first
            file_ext = Path(file_path).suffix.lower()
            
            if file_ext == '.pdf':
                text = self._extract_text_from_pdf_simple(file_path)
            elif file_ext in ['.docx', '.doc']:
                text = self._extract_text_from_docx_simple(file_path)
            else:
                text = ""
            
            return self._fallback_text_extraction(text)
            
        except Exception:
            return CVData()
    
    def _fallback_text_extraction(self, text: str) -> CVData:
        """Fallback extraction using regex patterns"""
        if not text:
            return CVData()
        
        # Extract basic info using regex
        email = self._extract_email(text)
        phone = self._extract_phone(text)
        skills = self.skill_manager.extract_skills_from_text(text)
        
        return CVData(
            email=email,
            phone=phone,
            technical_skills=skills
        )
    
    def _extract_text_from_pdf_simple(self, pdf_path: str) -> str:
        """Simple PDF text extraction"""
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
    
    def _extract_text_from_docx_simple(self, docx_path: str) -> str:
        """Simple DOCX text extraction"""
        doc = Document(docx_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    
    def _extract_email(self, text: str) -> Optional[str]:
        """Extract email using regex"""
        pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        matches = re.findall(pattern, text)
        return matches[0] if matches else None
    
    def _extract_phone(self, text: str) -> Optional[str]:
        """Extract phone using regex"""
        patterns = [
            r'(?:\+84|0)\s*[1-9]\d{8,9}',
            r'\(\d{2,3}\)\s*\d{3,4}[-\s]*\d{3,4}',
            r'\d{10,11}'
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text)
            if matches:
                return matches[0].replace(' ', '').replace('-', '')
        
        return None

def test_cv_extractor():
    """Test function"""
    extractor = CVExtractor()
    
    # Test with sample text
    sample_text = """
    Nguyễn Văn A
    Email: nguyenvana@email.com
    Phone: 0123456789
    
    Kỹ năng: Python, Django, ReactJS, MySQL, Docker
    
    Kinh nghiệm: 3 năm làm việc tại ABC Company
    """
    
    try:
        result = extractor.extract_from_text(sample_text)
        print("Extracted CV Data:")
        print(result.model_dump_json(indent=2))
    except Exception as e:
        print(f"Test error: {e}")

if __name__ == "__main__":
    test_cv_extractor() 